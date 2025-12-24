from __future__ import annotations

import csv
from pathlib import Path

import pytest

from amprenta_rag.extraction.parsers import get_parser, parse_csv, parse_docx


def test_parse_csv(tmp_path: Path) -> None:
    p = tmp_path / "demo.csv"
    with p.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["a", "b"])
        w.writerow([1, 2])
        w.writerow([3, 4])

    out = parse_csv(str(p))
    assert out["columns"] == ["a", "b"]
    assert out["row_count"] == 2
    assert isinstance(out["sample"], list)
    assert len(out["sample"]) == 2


def test_parse_docx(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document  # type: ignore

    p = tmp_path / "demo.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "col1"
    t.cell(0, 1).text = "col2"
    t.cell(1, 0).text = "v1"
    t.cell(1, 1).text = "v2"
    doc.save(str(p))

    out = parse_docx(str(p))
    assert "Hello world" in out["text"]
    assert isinstance(out["tables"], list)
    assert out["tables"]


def test_get_parser_csv() -> None:
    fn = get_parser("x.csv")
    assert fn is parse_csv


def test_get_parser_unsupported() -> None:
    with pytest.raises(ValueError):
        get_parser("x.unsupported")


